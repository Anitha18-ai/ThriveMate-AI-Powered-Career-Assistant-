import os
import io
import logging
import PyPDF2
from werkzeug.datastructures import FileStorage
import docx

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def extract_text_from_file(file: FileStorage) -> str:
    """
    Extract text from uploaded resume file (PDF or DOC/DOCX)
    
    Args:
        file: Uploaded file object
        
    Returns:
        Extracted text content as string
    """
    filename = file.filename
    file_extension = os.path.splitext(filename)[1].lower()
    
    try:
        # Process based on file extension
        if file_extension == '.pdf':
            return extract_text_from_pdf(file)
        elif file_extension in ['.doc', '.docx']:
            return extract_text_from_docx(file)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return ""
    except Exception as e:
        logger.exception(f"Error extracting text from file: {str(e)}")
        return ""

def extract_text_from_pdf(file: FileStorage) -> str:
    """
    Extract text from PDF file
    
    Args:
        file: Uploaded PDF file
        
    Returns:
        Extracted text content
    """
    text = ""
    
    try:
        # Create a file-like object from the file data
        file_stream = io.BytesIO(file.read())
        
        # Reset file pointer to beginning for PyPDF2
        file.seek(0)
        
        # Open the PDF file
        pdf_reader = PyPDF2.PdfReader(file_stream)
        
        # Extract text from each page
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
            
        return text
        
    except Exception as e:
        logger.exception(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file: FileStorage) -> str:
    """
    Extract text from DOCX file
    
    Args:
        file: Uploaded DOCX file
        
    Returns:
        Extracted text content
    """
    text = ""
    
    try:
        # Create a file-like object from the file data
        file_stream = io.BytesIO(file.read())
        
        # Reset file pointer to beginning
        file.seek(0)
        
        # Open the DOCX file
        doc = docx.Document(file_stream)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
        
    except Exception as e:
        logger.exception(f"Error extracting text from DOCX: {str(e)}")
        return ""

def clean_resume_text(text: str) -> str:
    """
    Clean up extracted text for better processing
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    text = ' '.join(text.split())
    
    # Replace multiple newlines with a single newline
    import re
    text = re.sub(r'\n+', '\n', text)
    
    # Remove special characters that might interfere with analysis
    text = re.sub(r'[^\w\s\n.,;:\-()]', '', text)
    
    return text
